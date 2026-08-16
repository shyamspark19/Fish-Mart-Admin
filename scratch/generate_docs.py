import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def create_documentation():
    doc = docx.Document()

    # Document Title
    title = doc.add_heading('Fish Mart — Complete Application Codebase & Function Documentation', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.color.rgb = RGBColor(7, 15, 30)
    title.style.font.name = 'Arial'

    subtitle = doc.add_paragraph('Detailed Technical Architecture, Code File Explanations, State Management, and Function Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Executive Summary
    h1 = doc.add_heading('1. Executive Overview & Technology Architecture', level=1)
    h1.style.font.color.rgb = RGBColor(2, 132, 199)

    p1 = doc.add_paragraph(
        "Fish Mart is a full-stack, enterprise-grade E-Commerce platform built for fresh seafood and fish delivery. "
        "The system features direct coastline catch sourcing, custom cutting preferences (Steak Cut, Curry Cut, Boneless Cubes), "
        "net vs gross weight pricing, delivery location selection powered by Google Maps API, a multi-tab Payment Gateway, "
        "and a dual-role architectural model featuring distinct visual design systems:"
    )

    bullet1 = doc.add_paragraph(style='List Bullet')
    r1 = bullet1.add_run('Customer Storefront (Blue Ocean Theme): ')
    r1.bold = True
    bullet1.add_run('Deep marine navy palette (#070F1E), electric cyan accents (#06B6D4), right-to-left sliding banner carousel, interactive cart drawer (visible after login), and 90-minute live order tracking.')

    bullet2 = doc.add_paragraph(style='List Bullet')
    r2 = bullet2.add_run('Admin Control Portal (Sunset Theme): ')
    r2.bold = True
    bullet2.add_run('Warm sunset orange palette (#F97316), amber highlights (#F59E0B), complete Read/Write product CRUD management, image picture URL live preview, stock modifiers (+5 / -5), customer order status manager, and Google Maps fulfillment hub detector.')

    # Tech stack table
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ['Layer / Component', 'Technologies Used']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], '0F172A')

    data = [
        ('Frontend Framework', 'React 18, Vite, TypeScript, Tailwind CSS, React Router v6'),
        ('Backend Framework', 'Node.js, Express.js, TypeScript, ts-node-dev'),
        ('Database & Fallback', 'MongoDB, Mongoose ORM, MongoMemoryServer (In-memory fallback)'),
        ('Authentication & Security', 'JSON Web Tokens (JWT), bcrypt password hashing, Role-Based Authorization'),
        ('Maps & External APIs', 'Google Maps iframe location embed, Geolocation API')
    ]

    for row_idx, (layer, tech) in enumerate(data, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        if row_idx % 2 == 0:
            set_cell_background(row_cells[0], 'F1F5F9')
            set_cell_background(row_cells[1], 'F1F5F9')

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Backend Files Section
    h2 = doc.add_heading('2. Backend Architecture & Server Files (server/src)', level=1)
    h2.style.font.color.rgb = RGBColor(2, 132, 199)

    # server.ts
    doc.add_heading('server/src/server.ts — Express Server Initialization', level=2)
    doc.add_paragraph(
        "Purpose: Entry point for the Node.js backend server.\n"
        "• Express App Configuration: Initializes middleware including cors() for cross-origin requests and express.json() for JSON body parsing.\n"
        "• Route Registration: Maps /api/auth, /api/products, /api/orders, and /api/admin/products to their respective route modules.\n"
        "• Database Connection Trigger: Calls connectDB() and auto-triggers seedInitialData() on startup to ensure initial catalog products exist."
    )

    # db.ts
    doc.add_heading('server/src/config/db.ts — Database Connection Manager', level=2)
    doc.add_paragraph(
        "Purpose: Manages MongoDB connection with automated in-memory fallback.\n"
        "• Dual Strategy: Attempts connection to local MongoDB (mongodb://localhost:27017/fishmart). If local MongoDB is unavailable, automatically initializes an in-memory MongoMemoryServer instance.\n"
        "• Zero-Downtime Reliability: Guarantees the application runs smoothly in any environment without requiring external database setup."
    )

    # seed.ts
    doc.add_heading('server/src/seed.ts — Initial Data Seeder', level=2)
    doc.add_paragraph(
        "Purpose: Seeds pre-populated Licious-style seafood products and default system users.\n"
        "• Catalog Products: Inserts 13 authentic seafood products with categories (Sea Fish, Freshwater Fish, Prawns & Shrimps, Crabs & Shellfish, Ready to Cook, Combo Packs), net/gross weights, cutting options, stock levels, and badges.\n"
        "• System Users: Creates default Admin (admin@fishmart.test / Admin123!) and Customer (customer@fishmart.test / Customer123!) accounts with bcrypt password hashes."
    )

    # Models
    doc.add_heading('server/src/models/ — Mongoose Schemas', level=2)
    doc.add_paragraph(
        "• Product.ts: Defines the schema for seafood items, including name, description, category, images array, weights array ({ label, price }), cuttingOptions array, stock number, badge, netWeight, grossWeight, rating, and reviewsCount.\n"
        "• User.ts: Defines user schema with name, email (unique), password hash, and role ('CUSTOMER' | 'ADMIN' | 'DELIVERY_PARTNER').\n"
        "• Order.ts: Defines order schema containing customer reference, items list, delivery address details, delivery slot, paymentMethod ('UPI' | 'CARD' | 'NETBANKING' | 'COD'), orderStatus ('PLACED' | 'CONFIRMED' | 'PREPARING' | 'OUT_FOR_DELIVERY' | 'DELIVERED' | 'CANCELLED'), subtotal, tax, and total amount."
    )

    # Controllers
    doc.add_heading('server/src/controllers/ — Request Handlers', level=2)
    doc.add_paragraph(
        "• authController.ts: Contains register() to create new user accounts and login() to verify credentials and return JWT bearer tokens.\n"
        "• productsController.ts: Contains getProducts() to list products with optional category filtering and getProductById().\n"
        "• adminProductController.ts: Admin-restricted handlers including createProduct(), updateProduct(), and deleteProduct().\n"
        "• ordersController.ts: Contains createOrder(), getUserOrders(), and updateOrderStatus() for tracking delivery status."
    )

    # Middlewares
    doc.add_heading('server/src/middleware/ — Security Middlewares', level=2)
    doc.add_paragraph(
        "• auth.ts (authenticate): Validates incoming Authorization: Bearer <token> headers and attaches the decoded user context to req.user.\n"
        "• roles.ts (authorizeRoles): Restricts route access to specified roles (e.g. 'ADMIN'). Returns 403 Forbidden if user lacks required role."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. Frontend Files Section
    h3 = doc.add_heading('3. Frontend Architecture & React Files (client/src)', level=1)
    h3.style.font.color.rgb = RGBColor(2, 132, 199)

    # Navbar
    doc.add_heading('client/src/components/Navbar.tsx — Application Navigation Bar', level=2)
    doc.add_paragraph(
        "Purpose: Sticky top navigation bar supporting dual-role states and clean auth layouts.\n"
        "• Route-Aware Minimal Header: On /login and /register pages, renders only the BrandLogo and a clean Sign In / Register link ('just app and signin is enough').\n"
        "• Changeable Location Trigger: Displays current delivery area (📍 Delivering to: T. Nagar, Chennai) and opens LocationModal on click.\n"
        "• Login-Gated Cart Badge: Displays Cart icon with live item counter badge ONLY when user is authenticated.\n"
        "• Role-Aware Navigation: Renders '👑 Admin Hub' button when logged in as ADMIN."
    )

    # BrandLogo
    doc.add_heading('client/src/components/BrandLogo.tsx — Brand Logo Component', level=2)
    doc.add_paragraph(
        "Purpose: Distinctive application logo and brand typography.\n"
        "• Image Rendering: Displays exact brand logo image with background contrast glassmorphic pill and cyan drop-shadow filter.\n"
        "• Fallback Badge & Typography: Includes onError fallback state rendering a high-visibility badge paired with Righteous / Syne Google Font typography ('FISH MART')."
    )

    # LocationModal
    doc.add_heading('client/src/components/LocationModal.tsx — Google Maps Location Picker', level=2)
    doc.add_paragraph(
        "Purpose: Interactive delivery location modal.\n"
        "• Google Maps Embed: Embedded iframe preview (https://maps.google.com/maps?q=lat,lng&output=embed) showing latitude/longitude pin.\n"
        "• GPS Auto-Detection: One-click button utilizing browser navigator.geolocation to detect current coordinates.\n"
        "• Preset Area Pills: Quick-select pills for popular Chennai delivery hubs (T. Nagar, Velachery, Anna Nagar, Adyar, Besant Nagar, OMR)."
    )

    # CartDrawer
    doc.add_heading('client/src/components/CartDrawer.tsx — Slide-Over Cart Drawer', level=2)
    doc.add_paragraph(
        "Purpose: Interactive slide-over cart drawer.\n"
        "• Quantity Controls: Adjust item quantities (+) / (-), remove items, select cut preferences (Steak Cut, Curry Cut, Boneless Cubes), and view weight variants.\n"
        "• Promo Coupons: Apply coupon codes OCEAN100 (₹100 off) or FRESH50 (₹50 off).\n"
        "• Live Totals Breakdown: Displays Subtotal, 5% Tax/Packing, Delivery Fee (Free above ₹499), and Total Payable with a direct Checkout CTA."
    )

    # Home.tsx
    doc.add_heading('client/src/pages/Home.tsx — Customer Storefront Page', level=2)
    doc.add_paragraph(
        "Purpose: Main customer landing and product catalog page.\n"
        "• Sliding Banner Carousel: Auto-playing 4-second right-to-left sliding hero banner showcasing fresh catch updates, jumbo prawns, and weekend promotions with slide indicator dots and arrow controls.\n"
        "• Search & Sort Toolbar: Search input for fish names and price sorting dropdown (Featured, Price: Low to High, Price: High to Low, Rating).\n"
        "• Category Filter Pills: Filter products by Sea Fish, Freshwater Fish, Prawns, Crabs, Ready to Cook, and Combo Packs.\n"
        "• Product Cards: Displays net vs gross weights, cut selectors, weight dropdowns, stock status, ratings, and Add to Cart handlers."
    )

    # AdminDashboard.tsx
    doc.add_heading('client/src/pages/AdminDashboard.tsx — Sunset Theme Admin Control Hub', level=2)
    doc.add_paragraph(
        "Purpose: Sunset Theme Admin Portal for complete inventory & order management.\n"
        "• Product Inventory Table (Read / Write): List all products with thumbnail images, prices, categories, stock controls (-5 / +5 buttons), edit modal triggers, and delete buttons.\n"
        "• Add / Edit Product Modal: Form to configure product name, category, price, stock, weight specs, cut preferences, and Image Picture URL with live picture preview.\n"
        "• Customer Orders Manager: View customer orders and update status dropdowns (PLACED → CONFIRMED → PREPARING → OUT_FOR_DELIVERY → DELIVERED → CANCELLED).\n"
        "• Google Maps Hub Detector: Interactive map showing central fulfillment hub coordinates and coverage area."
    )

    # Checkout.tsx
    doc.add_heading('client/src/pages/Checkout.tsx — Payment Gateway & Order Page', level=2)
    doc.add_paragraph(
        "Purpose: Complete checkout interface supporting 4 secure payment gateways.\n"
        "• Address & Slot Selection: Confirm recipient name, phone, delivery pin, and delivery time slot (ASAP 90 Mins, Evening, Tomorrow).\n"
        "• Payment Tabs: 1) UPI / QR Code (scan QR code / VPA ID), 2) Credit / Debit Card (interactive live card graphic preview with card number, expiry, CVV), 3) Net Banking (HDFC, ICICI, SBI, Axis, Kotak), 4) Cash on Delivery.\n"
        "• Simulated Payment Verification: Multi-stage loader overlay simulating 3D secure gateway authorization steps before finalizing order."
    )

    # OrderSuccess.tsx
    doc.add_heading('client/src/pages/OrderSuccess.tsx — Order Confirmation & Live Tracker', level=2)
    doc.add_paragraph(
        "Purpose: Order confirmation receipt page.\n"
        "• Celebration Card: Order reference number (FM...) and payment confirmation badge.\n"
        "• Live 4-Stage Tracker: Visual progress timeline (1. Placed → 2. Confirmed → 3. Out for Delivery → 4. Delivered) with 90-minute ETA timer.\n"
        "• Order Summary: Itemized receipt breakdown and delivery address confirmation."
    )

    # Login.tsx & Register.tsx
    doc.add_heading('client/src/pages/Login.tsx & Register.tsx — Role Authentication Pages', level=2)
    doc.add_paragraph(
        "Purpose: User login and registration with explicit Role Selection.\n"
        "• Role Selector Bar: Toggle between 👤 Customer and 👑 Admin Portal before logging in.\n"
        "• Role-Based Redirects: Logging in as Admin redirects directly to /admin (Sunset Theme Portal), while logging in as Customer redirects to / (Customer Storefront)."
    )

    # adminService.ts
    doc.add_heading('client/src/services/adminService.ts — Standalone Admin Helper Service', level=2)
    doc.add_paragraph(
        "Purpose: Standalone service module containing modular helper functions for admin CRUD operations:\n"
        "• updateProductStock(productId, newStock): Sends PUT request to /api/admin/products/:id with updated stock.\n"
        "• updateOrderStatus(orderId, status): Sends PUT request to /api/orders/:id/status with updated status.\n"
        "• deleteAdminProduct(productId): Sends DELETE request to remove product.\n"
        "• createAdminProduct(payload): Sends POST request to create new product.\n"
        "• editAdminProduct(productId, payload): Sends PUT request to modify product details."
    )

    # Save document
    filename = "d:/Esperia/3 Months Intern/Fish_Mart_Complete_Codebase_Documentation.docx"
    doc.save(filename)
    print(f"Documentation generated successfully at: {filename}")

if __name__ == '__main__':
    create_documentation()
